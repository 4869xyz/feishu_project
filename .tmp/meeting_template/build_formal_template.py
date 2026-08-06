from __future__ import annotations

import hashlib
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docxtpl import DocxTemplate


REFERENCE_SHA256 = "d3e1673e4bef5a60f1aca9600680d6d11d0c08f330518dffaf5e5a5269292666"
EXPECTED_KEYS = {
    "general_manager",
    "wu_aoxiang",
    "yang_yilin",
    "ye_mengzhen",
    "zhang_feilong",
    "lin_baili",
    "liang_jialong",
    "liu_hanwen",
    "xu_xinxin",
    "gao_canjian",
    "liu_jindi",
    "ma_gengbin",
    "zhang_chunwei",
}

DEPARTMENTS = (
    ("一、总经理", (("总经理", "general_manager"),)),
    ("二、商务部-销售组", (("1. 吴傲翔", "wu_aoxiang"), ("2. 杨意林", "yang_yilin"))),
    ("三、人事行政", (("叶梦真", "ye_mengzhen"),)),
    ("四、中台", (("张飞龙", "zhang_feilong"),)),
    ("五、运营推广组", (("林柏丽", "lin_baili"),)),
    ("六、后期部", (("1. 梁家龙", "liang_jialong"), ("2. 刘瀚文", "liu_hanwen"))),
    ("七、美术部", (("1. 徐兴鑫", "xu_xinxin"), ("2. 高灿健", "gao_canjian"), ("3. 刘锦娣", "liu_jindi"))),
    ("八、程序部", (("1. 马耿宾", "ma_gengbin"), ("2. 张春威", "zhang_chunwei"))),
)


def set_run_font(run, size: float, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "宋体"
    run.font.size = Pt(size)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(16)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run("周例会纪要"), 18, bold=True)


def add_department(document: Document, title: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(title), 14, bold=True)


def add_person(document: Document, label: str, key: str) -> None:
    name = document.add_paragraph()
    name.paragraph_format.space_before = Pt(4)
    name.paragraph_format.space_after = Pt(3)
    name.paragraph_format.keep_with_next = True
    set_run_font(name.add_run(label), 10.5, bold=True)

    slot = document.add_paragraph()
    slot.paragraph_format.space_after = Pt(8)
    slot.paragraph_format.line_spacing = 1.25
    set_run_font(slot.add_run("{{ " + key + " }}"), 10.5)


def main(reference_text: str, output_text: str) -> None:
    reference = Path(reference_text)
    output = Path(output_text)
    if hashlib.sha256(reference.read_bytes()).hexdigest() != REFERENCE_SHA256:
        raise RuntimeError("初始模板已变更，请重新检查后再生成正式模板。")

    document = Document(reference)
    section = document.sections[0]
    reference_geometry = (
        section.page_width,
        section.page_height,
        section.top_margin,
        section.bottom_margin,
        section.left_margin,
        section.right_margin,
    )
    clear_body(document)
    add_title(document)
    for department, people in DEPARTMENTS:
        add_department(document, department)
        for label, key in people:
            add_person(document, label, key)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)

    # Structural QA: the bot must recognize every placeholder and the source page setup must survive.
    template = DocxTemplate(output)
    found_keys = set(template.get_undeclared_template_variables())
    if found_keys != EXPECTED_KEYS:
        raise RuntimeError(f"占位符不完整或包含意外变量：{sorted(found_keys)}")
    verified = Document(output)
    result_section = verified.sections[0]
    result_geometry = (
        result_section.page_width,
        result_section.page_height,
        result_section.top_margin,
        result_section.bottom_margin,
        result_section.left_margin,
        result_section.right_margin,
    )
    if result_geometry != reference_geometry:
        raise RuntimeError("正式模板的页面几何与初始模板不一致。")
    print(f"Created: {output.resolve()}")
    print("Placeholders:", ", ".join(sorted(found_keys)))


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
