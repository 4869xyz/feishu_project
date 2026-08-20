from __future__ import annotations

import hashlib
import json
import sys
from pathlib import Path

from docx import Document


def paragraph_data(paragraph):
    return {
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else None,
        "alignment": str(paragraph.alignment),
        "runs": [
            {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "font": run.font.name,
                "size_pt": run.font.size.pt if run.font.size else None,
            }
            for run in paragraph.runs
        ],
    }


def main(path_text: str) -> None:
    path = Path(path_text)
    document = Document(path)
    data = {
        "reference": str(path.resolve()),
        "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
        "sections": [
            {
                "page_width": section.page_width,
                "page_height": section.page_height,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
                "header_distance": section.header_distance,
                "footer_distance": section.footer_distance,
            }
            for section in document.sections
        ],
        "paragraphs": [paragraph_data(p) for p in document.paragraphs],
        "tables": [
            [[paragraph_data(p) for p in cell.paragraphs] for cell in row.cells]
            for table in document.tables
            for row in table.rows
        ],
        "headers": [
            [paragraph_data(p) for p in section.header.paragraphs]
            for section in document.sections
        ],
        "footers": [
            [paragraph_data(p) for p in section.footer.paragraphs]
            for section in document.sections
        ],
        "inline_shapes": len(document.inline_shapes),
        "styles": [style.name for style in document.styles],
    }
    print(json.dumps(data, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main(sys.argv[1])
