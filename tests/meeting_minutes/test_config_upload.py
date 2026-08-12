"""Admin private-chat upload of people YAML and Word templates."""

from __future__ import annotations

import asyncio
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document

from meeting_minutes_bot.attachments import (
    AttachmentProcessingError,
    AttachmentResource,
    ExtractedAttachment,
    validate_resource_type,
)
from meeting_minutes_bot.config_update import (
    CONFIG_BACKUP_DIRNAME,
    classify_admin_config_upload,
)
from meeting_minutes_bot.document import MinutesDocumentRenderer
from meeting_minutes_bot.people import PeopleStore
from meeting_minutes_bot.repository import MeetingRepository
from meeting_minutes_bot.service import (
    VALIDATE_COMMAND,
    MeetingMinutesService,
)
from tests.meeting_minutes.helpers import TEST_TEMPLATE
from tests.meeting_minutes.test_people_reload import (
    INITIAL_PEOPLE_YAML,
    REBOUND_PEOPLE_YAML,
)


NOW = datetime(2026, 8, 5, 12, 0)


async def build_upload_service(
    root: Path,
) -> tuple[MeetingMinutesService, MeetingRepository, PeopleStore, Path, Path]:
    """Build a service with writable people YAML and a copy of the test template."""

    data_dir = root / "data"
    data_dir.mkdir(parents=True, exist_ok=True)
    config_path = root / "people.yaml"
    config_path.write_text(INITIAL_PEOPLE_YAML, encoding="utf-8")
    template_path = root / TEST_TEMPLATE.name
    shutil.copy2(TEST_TEMPLATE, template_path)

    store = PeopleStore.from_path(config_path)
    repository = MeetingRepository(
        f"sqlite+aiosqlite:///{(root / 'meeting.db').as_posix()}"
    )
    await repository.initialize()
    renderer = MinutesDocumentRenderer(
        template_path=template_path,
        output_dir=root / "output",
        people=store,
        data_dir=data_dir,
    )
    service = MeetingMinutesService(
        repository=repository,
        people=store,
        renderer=renderer,
        timezone="Asia/Shanghai",
        max_text_length=100,
        data_dir=data_dir,
    )
    return service, repository, store, config_path, template_path


def test_classify_admin_config_upload_rules() -> None:
    template = Path("templates") / "周例会纪要正式模板.docx"
    assert (
        classify_admin_config_upload(
            file_name="people.yaml", template_path=template, is_admin=True
        )
        == "people"
    )
    assert (
        classify_admin_config_upload(
            file_name="周例会纪要正式模板.docx",
            template_path=template,
            is_admin=True,
        )
        == "template"
    )
    assert (
        classify_admin_config_upload(
            file_name="我的周报.docx", template_path=template, is_admin=True
        )
        is None
    )
    assert (
        classify_admin_config_upload(
            file_name="people.yaml", template_path=template, is_admin=False
        )
        is None
    )


def test_validate_resource_type_allows_yaml_only_for_admin_config() -> None:
    resource = AttachmentResource("file", "file_yaml", "people.yaml")
    assert validate_resource_type(resource, allow_admin_config=True) == "yaml"
    try:
        validate_resource_type(resource, allow_admin_config=False)
        raise AssertionError("expected rejection")
    except AttachmentProcessingError as exc:
        assert "不支持该附件格式" in str(exc)


def test_people_yaml_upload_applies_and_backs_up(project_tmp_dir: Path) -> None:
    async def scenario() -> None:
        service, repository, store, config_path, _ = await build_upload_service(
            project_tmp_dir
        )
        try:
            before = config_path.read_bytes()
            upload = project_tmp_dir / "upload_people.yaml"
            upload.write_text(REBOUND_PEOPLE_YAML, encoding="utf-8")
            result = await service.handle_config_upload(
                sender_open_id="ou_admin",
                file_name="people.yaml",
                source_path=upload,
            )
            assert "人员配置已通过上传更新并生效" in result.text
            assert "启用人数：2" in result.text
            assert config_path.read_text(encoding="utf-8") == REBOUND_PEOPLE_YAML
            backups = list((service.data_dir / CONFIG_BACKUP_DIRNAME).glob("*.bak"))
            assert len(backups) == 1
            assert backups[0].read_bytes() == before
            accepted = await service.handle_text(
                message_id="om_u1",
                sender_open_id="ou_new",
                text="新人内容",
                received_at=NOW,
            )
            assert "已收到" in accepted.text
            assert store.directory.find("ou_wu") is None
        finally:
            await repository.close()

    asyncio.run(scenario())


def test_people_yaml_upload_keeps_old_on_invalid(project_tmp_dir: Path) -> None:
    async def scenario() -> None:
        service, repository, store, config_path, _ = await build_upload_service(
            project_tmp_dir
        )
        try:
            before_dir = store.directory
            before_bytes = config_path.read_bytes()
            upload = project_tmp_dir / "bad.yaml"
            upload.write_text("people: [broken]", encoding="utf-8")
            result = await service.handle_config_upload(
                sender_open_id="ou_admin",
                file_name="people.yaml",
                source_path=upload,
            )
            assert "配置更新失败" in result.text
            assert store.directory is before_dir
            assert config_path.read_bytes() == before_bytes
            assert not (service.data_dir / CONFIG_BACKUP_DIRNAME).exists() or not list(
                (service.data_dir / CONFIG_BACKUP_DIRNAME).glob("*.bak")
            )
        finally:
            await repository.close()

    asyncio.run(scenario())


def test_config_upload_requires_admin(project_tmp_dir: Path) -> None:
    async def scenario() -> None:
        service, repository, _, _, _ = await build_upload_service(project_tmp_dir)
        try:
            upload = project_tmp_dir / "people.yaml"
            upload.write_text(REBOUND_PEOPLE_YAML, encoding="utf-8")
            result = await service.handle_config_upload(
                sender_open_id="ou_wu",
                file_name="people.yaml",
                source_path=upload,
            )
            assert "没有上传配置文件的权限" in result.text
        finally:
            await repository.close()

    asyncio.run(scenario())


def test_template_upload_same_name_applies(project_tmp_dir: Path) -> None:
    async def scenario() -> None:
        service, repository, _, _, template_path = await build_upload_service(
            project_tmp_dir
        )
        try:
            before = template_path.read_bytes()
            upload = project_tmp_dir / "new_template.docx"
            shutil.copy2(TEST_TEMPLATE, upload)
            # Touch content so bytes differ while placeholders stay valid.
            document = Document(str(upload))
            document.add_paragraph("template-upload-marker")
            document.save(str(upload))
            assert upload.read_bytes() != before

            result = await service.handle_config_upload(
                sender_open_id="ou_admin",
                file_name=template_path.name,
                source_path=upload,
            )
            assert "Word 模板已通过上传更新并生效" in result.text
            assert template_path.read_bytes() == upload.read_bytes()
            backups = list((service.data_dir / CONFIG_BACKUP_DIRNAME).glob("*.bak"))
            assert len(backups) == 1
            assert backups[0].read_bytes() == before
        finally:
            await repository.close()

    asyncio.run(scenario())


def test_template_upload_missing_placeholder_keeps_old(
    project_tmp_dir: Path,
) -> None:
    async def scenario() -> None:
        service, repository, _, _, template_path = await build_upload_service(
            project_tmp_dir
        )
        try:
            before = template_path.read_bytes()
            upload = project_tmp_dir / "empty.docx"
            Document().save(str(upload))
            result = await service.handle_config_upload(
                sender_open_id="ou_admin",
                file_name=template_path.name,
                source_path=upload,
            )
            assert "配置更新失败" in result.text
            assert template_path.read_bytes() == before
        finally:
            await repository.close()

    asyncio.run(scenario())


def test_admin_non_template_docx_still_submits(project_tmp_dir: Path) -> None:
    async def scenario() -> None:
        # Admin who is also an enabled person can submit a differently named DOCX.
        yaml_text = (
            "people:\n"
            "  ou_admin:\n"
            "    name: 管理员\n"
            "    department: 商务部-销售组\n"
            "    template_key: wu_aoxiang\n"
            "    section_order: 2\n"
            "    sort_order: 1\n"
            "admins:\n"
            "  - ou_admin\n"
        )
        service, repository, store, config_path, template_path = (
            await build_upload_service(project_tmp_dir)
        )
        try:
            before_template = template_path.read_bytes()
            config_path.write_text(yaml_text, encoding="utf-8")
            store.replace(store.load_candidate())
            assert service.classify_config_upload("ou_admin", "我的周报.docx") is None

            source = project_tmp_dir / "我的周报.docx"
            doc = Document()
            doc.add_paragraph("admin weekly note")
            doc.save(str(source))
            attachment = ExtractedAttachment(
                file_name="我的周报.docx",
                message_type="docx",
                raw_content="admin weekly note",
                parsed_content="admin weekly note",
                recognition_method="docx",
                source_path=source,
            )
            result = await service.handle_attachment(
                message_id="om_u_docx",
                sender_open_id="ou_admin",
                attachment=attachment,
                received_at=NOW,
            )
            assert "已识别并追加" in result.text or "已收到" in result.text
            assert template_path.read_bytes() == before_template
            backup_dir = service.data_dir / CONFIG_BACKUP_DIRNAME
            assert not backup_dir.exists() or not list(backup_dir.glob("*.bak"))
        finally:
            await repository.close()

    asyncio.run(scenario())


def test_validate_config_command(project_tmp_dir: Path) -> None:
    async def scenario() -> None:
        service, repository, _, _, _ = await build_upload_service(project_tmp_dir)
        try:
            denied = await service.handle_text(
                message_id="om_v0",
                sender_open_id="ou_wu",
                text=VALIDATE_COMMAND,
                received_at=NOW,
            )
            assert "没有执行" in denied.text
            result = await service.handle_text(
                message_id="om_v1",
                sender_open_id="ou_admin",
                text=VALIDATE_COMMAND,
                received_at=NOW,
            )
            assert "当前配置体检" in result.text
            assert "状态：通过" in result.text
        finally:
            await repository.close()

    asyncio.run(scenario())


def test_admin_warning_when_removed_from_admins(project_tmp_dir: Path) -> None:
    async def scenario() -> None:
        service, repository, _, _, _ = await build_upload_service(project_tmp_dir)
        try:
            yaml_text = REBOUND_PEOPLE_YAML.replace(
                "admins:\n  - ou_admin", "admins:\n  - ou_other"
            )
            upload = project_tmp_dir / "no_self.yaml"
            upload.write_text(yaml_text, encoding="utf-8")
            result = await service.handle_config_upload(
                sender_open_id="ou_admin",
                file_name="people.yaml",
                source_path=upload,
            )
            assert "人员配置已通过上传更新并生效" in result.text
            assert "未包含你的 open_id" in result.text
            denied = await service.handle_text(
                message_id="om_w1",
                sender_open_id="ou_admin",
                text=VALIDATE_COMMAND,
                received_at=NOW,
            )
            assert "没有执行" in denied.text
        finally:
            await repository.close()

    asyncio.run(scenario())
