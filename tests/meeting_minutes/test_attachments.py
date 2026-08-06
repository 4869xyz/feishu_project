from __future__ import annotations

import asyncio
from pathlib import Path
from types import SimpleNamespace

import pymupdf
import pytest
from docx import Document
from PIL import Image

from meeting_minutes_bot.attachments import (
    AttachmentProcessingError,
    AttachmentProcessor,
    AttachmentResource,
    ExtractedAttachment,
    message_attachment_resource,
    validate_resource_type,
)
from meeting_minutes_bot.listener import handle_message
from meeting_minutes_bot.period import meeting_period
from tests.meeting_minutes.helpers import build_service


class FakeOcr:
    def __init__(self, text: str = "识别出的图片文字") -> None:
        self.text = text
        self.paths: list[Path] = []

    def recognize(self, path: Path) -> str:
        self.paths.append(path)
        return self.text


def processor(*, ocr: FakeOcr | None = None, max_pages: int = 50) -> AttachmentProcessor:
    return AttachmentProcessor(
        ocr=ocr or FakeOcr(),
        max_bytes=20 * 1024 * 1024,
        max_pdf_pages=max_pages,
    )


def extract(
    worker: AttachmentProcessor, path: Path, resource: AttachmentResource
) -> ExtractedAttachment:
    return asyncio.run(worker.extract(path, resource))


def test_resource_metadata_supports_one_image_or_file() -> None:
    resource = message_attachment_resource(
        SimpleNamespace(
            resources=[
                SimpleNamespace(type="file", file_key="file_1", file_name="纪要.md")
            ]
        )
    )
    assert resource == AttachmentResource("file", "file_1", "纪要.md")
    assert validate_resource_type(resource) == "markdown"

    with pytest.raises(AttachmentProcessingError, match="只支持一个"):
        message_attachment_resource(
            SimpleNamespace(
                resources=[
                    SimpleNamespace(type="image", file_key="img_1"),
                    SimpleNamespace(type="image", file_key="img_2"),
                ]
            )
        )

    with pytest.raises(AttachmentProcessingError, match="不支持"):
        validate_resource_type(AttachmentResource("file", "file_2", "旧文档.doc"))


def test_extracts_image_with_local_ocr(project_tmp_dir: Path) -> None:
    image_path = project_tmp_dir / "纪要.png"
    Image.new("RGB", (120, 80), "white").save(image_path)
    fake = FakeOcr("本周完成招聘\n下周整理档案")

    result = extract(
        processor(ocr=fake),
        image_path,
        AttachmentResource("image", "img_1", "纪要.png"),
    )

    assert result.message_type == "image"
    assert result.parsed_content == "本周完成招聘\n下周整理档案"
    assert result.recognition_method == "本地图片 OCR"
    assert fake.paths == [image_path.resolve()]


def test_rejects_damaged_and_oversized_pixel_images(
    project_tmp_dir: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    damaged = project_tmp_dir / "damaged.png"
    damaged.write_bytes(b"not-an-image")
    resource = AttachmentResource("file", "file_bad", damaged.name)
    with pytest.raises(AttachmentProcessingError, match="损坏"):
        extract(processor(), damaged, resource)

    large = project_tmp_dir / "large.png"
    Image.new("RGB", (100, 100), "white").save(large)
    monkeypatch.setattr("meeting_minutes_bot.attachments.MAX_IMAGE_PIXELS", 100)
    with pytest.raises(AttachmentProcessingError, match="像素尺寸超过限制"):
        extract(processor(), large, AttachmentResource("file", "file_large", large.name))


def test_extracts_docx_paragraphs_and_tables(project_tmp_dir: Path) -> None:
    path = project_tmp_dir / "纪要.docx"
    document = Document()
    document.add_paragraph("本周完成事项")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "负责人"
    table.cell(0, 1).text = "薛轶"
    document.save(path)

    result = extract(
        processor(), path, AttachmentResource("file", "file_docx", path.name)
    )

    assert result.message_type == "docx"
    assert result.parsed_content == "本周完成事项\n负责人 | 薛轶"


def test_extracts_markdown_as_plain_text(project_tmp_dir: Path) -> None:
    path = project_tmp_dir / "纪要.md"
    path.write_text(
        "# 本周完成\n\n- 招聘面试\n- [整理档案](https://example.com)\n\n```python\nprint('ok')\n```",
        encoding="utf-8-sig",
    )

    result = extract(
        processor(), path, AttachmentResource("file", "file_md", path.name)
    )

    assert result.message_type == "markdown"
    assert "本周完成" in result.parsed_content
    assert "- 招聘面试" in result.parsed_content
    assert "整理档案" in result.parsed_content
    assert "print('ok')" in result.parsed_content
    assert "https://example.com" not in result.parsed_content


def _make_pdf(path: Path, pages: tuple[str | None, ...], *, password: str = "") -> None:
    document = pymupdf.open()
    for text in pages:
        page = document.new_page()
        if text:
            page.insert_text((72, 72), text)
    if password:
        document.save(
            path,
            encryption=pymupdf.PDF_ENCRYPT_AES_256,
            owner_pw="owner-password",
            user_pw=password,
        )
    else:
        document.save(path)
    document.close()


def test_extracts_text_pdf_without_ocr(project_tmp_dir: Path) -> None:
    path = project_tmp_dir / "minutes.pdf"
    _make_pdf(path, ("First page", "Second page"))
    fake = FakeOcr("must not be used")

    result = extract(
        processor(ocr=fake), path, AttachmentResource("file", "file_pdf", path.name)
    )

    assert result.message_type == "pdf"
    assert "First page" in result.parsed_content
    assert "Second page" in result.parsed_content
    assert fake.paths == []


@pytest.mark.parametrize("pages", [(None,), ("Text page", None)])
def test_rejects_scanned_or_mixed_pdf_pages(
    project_tmp_dir: Path, pages: tuple[str | None, ...]
) -> None:
    path = project_tmp_dir / "scanned.pdf"
    _make_pdf(path, pages)
    with pytest.raises(AttachmentProcessingError, match="不支持扫描版或图片型 PDF"):
        extract(processor(), path, AttachmentResource("file", "file_pdf", path.name))


def test_rejects_encrypted_over_page_limit_and_damaged_pdfs(
    project_tmp_dir: Path,
) -> None:
    encrypted = project_tmp_dir / "encrypted.pdf"
    _make_pdf(encrypted, ("secret",), password="read-password")
    with pytest.raises(AttachmentProcessingError, match="加密"):
        extract(
            processor(),
            encrypted,
            AttachmentResource("file", "file_encrypted", encrypted.name),
        )

    too_long = project_tmp_dir / "too-long.pdf"
    _make_pdf(too_long, ("one", "two", "three"))
    with pytest.raises(AttachmentProcessingError, match="页数超过限制"):
        extract(
            processor(max_pages=2),
            too_long,
            AttachmentResource("file", "file_long", too_long.name),
        )

    damaged = project_tmp_dir / "damaged.pdf"
    damaged.write_bytes(b"not-a-pdf")
    with pytest.raises(AttachmentProcessingError, match="损坏"):
        extract(
            processor(),
            damaged,
            AttachmentResource("file", "file_bad", damaged.name),
        )


class FakeChannel:
    def __init__(self, cached_path: Path) -> None:
        self.cached_path = cached_path
        self.calls: list[tuple[str, dict, dict]] = []
        self.downloads = 0

    async def send(self, chat_id: str, content: dict, options: dict) -> None:
        self.calls.append((chat_id, content, options))

    async def resolve_resource_to_cache(self, **kwargs: object) -> object:
        self.downloads += 1
        return SimpleNamespace(decision="cached", path=self.cached_path)


def _file_message(sender: str, message_id: str = "om_file") -> object:
    return SimpleNamespace(
        chat_id="oc_chat",
        message_id=message_id,
        sender_id=sender,
        raw_content_type="file",
        resources=[
            SimpleNamespace(type="file", file_key="file_md", file_name="纪要.md")
        ],
    )


def test_listener_auto_appends_attachment_and_rejects_unbound_before_download(
    project_tmp_dir: Path,
) -> None:
    async def scenario() -> None:
        path = project_tmp_dir / "纪要.md"
        path.write_text("# 已完成\n\n- 招聘", encoding="utf-8")
        service, repository = await build_service(project_tmp_dir)
        channel = FakeChannel(path)
        worker = processor()
        try:
            await handle_message(channel, service, _file_message("ou_yang"), worker)
            assert channel.downloads == 1
            assert "正在本地识别" in channel.calls[0][1]["text"]
            assert "已识别并追加" in channel.calls[1][1]["text"]
            rows = await repository.submissions_for_person(
                period=meeting_period(None), open_id="ou_yang"
            )
            assert len(rows) == 1
            assert rows[0].message_type == "markdown"
            assert "已完成" in rows[0].parsed_content

            await handle_message(channel, service, _file_message("ou_yang"), worker)
            repeated = await repository.submissions_for_person(
                period=meeting_period(None), open_id="ou_yang"
            )
            assert len(repeated) == 1
            assert "已经处理" in channel.calls[-1][1]["text"]

            downloads_before_unbound = channel.downloads
            await handle_message(
                channel,
                service,
                _file_message("ou_unknown", message_id="om_unknown"),
                worker,
            )
            assert channel.downloads == downloads_before_unbound
            assert "尚未绑定" in channel.calls[-1][1]["text"]
        finally:
            await repository.close()

    asyncio.run(scenario())


def test_attachment_text_over_limit_is_rejected_without_truncation(
    project_tmp_dir: Path,
) -> None:
    async def scenario() -> None:
        service, repository = await build_service(project_tmp_dir)
        try:
            result = await service.handle_attachment(
                message_id="om_long_attachment",
                sender_open_id="ou_yang",
                attachment=ExtractedAttachment(
                    file_name="long.md",
                    message_type="markdown",
                    raw_content="x" * 101,
                    parsed_content="x" * 101,
                    recognition_method="Markdown 文本解析",
                ),
            )
            assert "超过 100 字限制" in result.text
            rows = await repository.submissions_for_person(
                period=meeting_period(None), open_id="ou_yang"
            )
            assert rows == ()
        finally:
            await repository.close()

    asyncio.run(scenario())
