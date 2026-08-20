"""DOCX submissions with tables/images are persisted and embedded into weekly DOCX."""

from __future__ import annotations

import asyncio
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image

from meeting_minutes_bot.attachments import AttachmentResource
from meeting_minutes_bot.docx_merge import compose_person_docx
from meeting_minutes_bot.service import GENERATE_COMMAND
from tests.meeting_minutes.helpers import build_service
from tests.meeting_minutes.test_attachments import FakeOcr, extract, processor


NOW = datetime(2026, 8, 5, 12, 0)


def _build_rich_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("本周完成事项")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "负责人"
    table.cell(0, 1).text = "薛轶"
    table.cell(1, 0).text = "进度"
    table.cell(1, 1).text = "已完成"
    image = path.with_suffix(".png")
    Image.new("RGB", (48, 24), color=(20, 120, 200)).save(image)
    document.add_paragraph().add_run().add_picture(str(image), width=Inches(1.2))
    document.save(path)


def test_extract_docx_accepts_image_only_document(project_tmp_dir: Path) -> None:
    path = project_tmp_dir / "only_image.docx"
    document = Document()
    image = project_tmp_dir / "dot.png"
    Image.new("RGB", (16, 16), color="green").save(image)
    document.add_paragraph().add_run().add_picture(str(image), width=Inches(0.5))
    document.save(path)

    result = extract(
        processor(ocr=FakeOcr("unused")),
        path,
        AttachmentResource("file", "file_img_docx", path.name),
    )

    assert result.message_type == "docx"
    assert result.has_embedded_media
    assert result.source_path == path.resolve()
    assert "表格或图片" in result.parsed_content or result.parsed_content


def test_extract_docx_keeps_table_summary_and_media_flag(project_tmp_dir: Path) -> None:
    path = project_tmp_dir / "rich.docx"
    _build_rich_docx(path)

    result = extract(
        processor(),
        path,
        AttachmentResource("file", "file_rich", path.name),
    )

    assert "本周完成事项" in result.parsed_content
    assert "负责人 | 薛轶" in result.parsed_content
    assert result.has_embedded_media
    assert "原样嵌入" in result.preview


def test_compose_person_docx_keeps_table_and_image(project_tmp_dir: Path) -> None:
    source = project_tmp_dir / "source.docx"
    _build_rich_docx(source)
    output = project_tmp_dir / "composed.docx"

    compose_person_docx(items=[("摘要", source)], output_path=output)
    document = Document(str(output))

    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 1).text.strip() == "薛轶"
    assert len(document.inline_shapes) >= 1
    image_rels = [
        rel
        for rel in document.part.rels.values()
        if "image" in str(getattr(rel, "reltype", "")).lower()
    ]
    assert image_rels


def test_generate_embeds_docx_table_and_image(project_tmp_dir: Path) -> None:
    async def scenario() -> None:
        service, repository = await build_service(project_tmp_dir)
        try:
            source = project_tmp_dir / "submit.docx"
            _build_rich_docx(source)
            attachment = await processor().extract(
                source,
                AttachmentResource("file", "file_submit", source.name),
            )
            accepted = await service.handle_attachment(
                message_id="om_docx_1",
                sender_open_id="ou_yang",
                attachment=attachment,
                received_at=NOW,
            )
            assert "已识别并追加" in accepted.text

            generated = await service.handle_text(
                message_id="om_gen_docx",
                sender_open_id="ou_admin",
                text=GENERATE_COMMAND,
                received_at=NOW,
            )
            assert generated.file_path is not None
            assert generated.file_path.is_file()

            output = Document(str(generated.file_path))
            assert any(table.cell(0, 1).text.strip() == "薛轶" for table in output.tables)
            image_rels = [
                rel
                for rel in output.part.rels.values()
                if "image" in str(getattr(rel, "reltype", "")).lower()
            ]
            assert image_rels

            persisted = list((project_tmp_dir / "data" / "submission_docs").rglob("*.docx"))
            assert persisted
        finally:
            await repository.close()

    asyncio.run(scenario())
