"""Build the deterministic DOCX fixture used by meeting-minutes tests."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PEOPLE = (
    ("一、总经理", (("总经理", "general_manager"),)),
    ("二、商务部-销售组", (("1、吴傲翔", "wu_aoxiang"), ("2、杨意林", "yang_yilin"))),
    ("三、人事行政", (("叶梦真", "ye_mengzhen"),)),
    ("四、中台", (("张飞龙", "zhang_feilong"),)),
    ("五、运营推广组", (("林柏丽", "lin_baili"),)),
    ("六、后期部", (("1、梁家龙", "liang_jialong"), ("2、刘瀚文", "liu_hanwen"))),
    ("七、美术部", (("1、徐兴鑫", "xu_xinxin"), ("2、高灿健", "gao_canjian"), ("3、刘锦娣", "liu_jindi"))),
    ("八、程序部", (("1、马耿宾", "ma_gengbin"), ("2、张春威", "zhang_chunwei"))),
)


def _font(run, name: str = "Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)


def build(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("周例会纪要")
    _font(run)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 77, 120)

    for heading, members in PEOPLE:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(heading)
        _font(run)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(46, 116, 181)
        for label, key in members:
            label_paragraph = document.add_paragraph()
            label_paragraph.paragraph_format.space_after = Pt(2)
            label_run = label_paragraph.add_run(label)
            _font(label_run)
            label_run.bold = True
            content_paragraph = document.add_paragraph()
            content_paragraph.paragraph_format.left_indent = Inches(0.25)
            content_run = content_paragraph.add_run("{{ " + key + " }}")
            _font(content_run)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


if __name__ == "__main__":
    build(Path(__file__).with_name("周例会纪要测试模板.docx"))
