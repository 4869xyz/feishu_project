"""Merge Word body content (including inline images) into a target document."""

from __future__ import annotations

from copy import deepcopy
from io import BytesIO
import logging
from pathlib import Path
import shutil

from docx import Document
from docx.oxml.ns import qn


LOGGER = logging.getLogger(__name__)
SUBMISSION_DOCS_DIRNAME = "submission_docs"


def submission_docs_root(data_dir: Path) -> Path:
    return Path(data_dir) / SUBMISSION_DOCS_DIRNAME


def persist_submission_docx(
    *,
    source: Path,
    data_dir: Path,
    period: str,
    message_id: str,
) -> str:
    """Copy a submitted DOCX into durable storage; return path relative to data_dir."""

    safe_id = "".join(
        char if char.isalnum() or char in {"-", "_"} else "_" for char in message_id
    )
    relative = Path(SUBMISSION_DOCS_DIRNAME) / period / f"{safe_id}.docx"
    target = Path(data_dir) / relative
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target)
    return relative.as_posix()


def _image_rids(element) -> list[str]:
    """Collect relationship ids referenced by drawing/blip elements."""

    found: list[str] = []
    for blip in element.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if rid:
            found.append(rid)
    for node in element.iter():
        for attr_name, attr_value in node.attrib.items():
            if attr_name.endswith("}embed") or attr_name.endswith("}id"):
                if isinstance(attr_value, str) and attr_value.startswith("rId"):
                    found.append(attr_value)
    # Preserve order while de-duplicating.
    return list(dict.fromkeys(found))


def _copy_image_relationship(target: Document, source: Document, rid: str) -> str | None:
    """Copy one image part from source into target and return the new rId."""

    try:
        rel = source.part.rels[rid]
    except KeyError:
        return None
    try:
        image_blob = rel.target_part.blob
        # Returns (rId, Image); we only need the new relationship id.
        new_rid, _image = target.part.get_or_add_image(BytesIO(image_blob))
        return new_rid
    except Exception:
        LOGGER.exception("复制 Word 图片关系到目标文档失败：rid=%s", rid)
        return None


def _clone_body_child(
    target: Document, source: Document, child, rid_map: dict[str, str]
):
    """Deep-copy one body child into ``target``, remapping image relationships."""

    cloned = deepcopy(child)
    for old_rid in _image_rids(cloned):
        if old_rid not in rid_map:
            mapped = _copy_image_relationship(target, source, old_rid)
            if mapped:
                rid_map[old_rid] = mapped
        new_rid = rid_map.get(old_rid)
        if not new_rid:
            continue
        for blip in cloned.iter(qn("a:blip")):
            if blip.get(qn("r:embed")) == old_rid:
                blip.set(qn("r:embed"), new_rid)
            if blip.get(qn("r:link")) == old_rid:
                blip.set(qn("r:link"), new_rid)
    return cloned


def append_docx_body(target: Document, source_path: Path) -> None:
    """Append body block-level elements from ``source_path`` into ``target``."""

    source = Document(str(source_path))
    rid_map: dict[str, str] = {}
    for child in list(source.element.body):
        if child.tag.endswith("}sectPr"):
            continue
        target.element.body.append(_clone_body_child(target, source, child, rid_map))


def insert_docx_body_before(
    target: Document, source_path: Path, anchor_element
) -> None:
    """Insert body elements from ``source_path`` immediately before ``anchor_element``."""

    source = Document(str(source_path))
    rid_map: dict[str, str] = {}
    for child in list(source.element.body):
        if child.tag.endswith("}sectPr"):
            continue
        cloned = _clone_body_child(target, source, child, rid_map)
        anchor_element.addprevious(cloned)


def compose_person_docx(
    *,
    items: list[tuple[str, Path | None]],
    output_path: Path,
) -> Path:
    """Build one DOCX for a person from ordered (summary, optional source) items."""

    document = Document()
    # Remove the default empty paragraph for a cleaner merge.
    if document.paragraphs and not document.paragraphs[0].text:
        element = document.paragraphs[0]._element
        element.getparent().remove(element)

    for index, (summary, source_path) in enumerate(items, 1):
        document.add_paragraph(f"{index}.")
        if source_path is not None and source_path.is_file():
            try:
                append_docx_body(document, source_path)
                continue
            except Exception:
                LOGGER.exception(
                    "合并源 Word 失败，回退文字摘要：%s", source_path
                )
        text = summary.strip() or "（无文字内容）"
        for line in text.splitlines() or ["（无文字内容）"]:
            document.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
